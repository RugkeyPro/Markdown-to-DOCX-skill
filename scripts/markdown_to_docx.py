#!/usr/bin/env python
"""Convert Markdown to DOCX with Pandoc, then apply the Chinese article style."""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT_CN = "SimSun"
FONT_CN_HEADING = "SimHei"
FONT_EN = "Times New Roman"
BLACK = RGBColor(0, 0, 0)


def run(cmd: list[str], *, check: bool = True) -> subprocess.CompletedProcess:
    return subprocess.run(cmd, check=check, text=True)


def find_pandoc() -> str | None:
    found = shutil.which("pandoc")
    if found:
        return found
    candidates = [
        Path(os.environ.get("LOCALAPPDATA", "")) / "Pandoc" / "pandoc.exe",
        Path(os.environ.get("ProgramFiles", "")) / "Pandoc" / "pandoc.exe",
        Path(os.environ.get("ProgramFiles(x86)", "")) / "Pandoc" / "pandoc.exe",
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def install_pandoc() -> str:
    if os.name == "nt":
        if shutil.which("winget"):
            run([
                "winget", "install", "--id", "JohnMacFarlane.Pandoc", "-e", "--silent",
                "--accept-package-agreements", "--accept-source-agreements",
            ])
        elif shutil.which("choco"):
            run(["choco", "install", "pandoc", "-y"])
        elif shutil.which("scoop"):
            run(["scoop", "install", "pandoc"])
        else:
            raise RuntimeError("Pandoc is missing and no supported Windows package manager was found.")
    elif sys.platform == "darwin":
        if not shutil.which("brew"):
            raise RuntimeError("Pandoc is missing and Homebrew was not found.")
        run(["brew", "install", "pandoc"])
    else:
        if shutil.which("apt-get"):
            run(["sudo", "apt-get", "update"])
            run(["sudo", "apt-get", "install", "-y", "pandoc"])
        elif shutil.which("dnf"):
            run(["sudo", "dnf", "install", "-y", "pandoc"])
        elif shutil.which("pacman"):
            run(["sudo", "pacman", "-S", "--noconfirm", "pandoc"])
        else:
            raise RuntimeError("Pandoc is missing and no supported Linux package manager was found.")
    pandoc = find_pandoc()
    if not pandoc:
        raise RuntimeError("Pandoc installation completed, but pandoc was not found on PATH.")
    return pandoc


def escape_yaml(value: str) -> str:
    return value.replace("\\", "\\\\").replace('"', '\\"')


def heading_text(line: str, level: int) -> str | None:
    match = re.match(rf"^#{{{level}}}\s+(.+?)\s*$", line)
    return match.group(1).strip() if match else None


def normalized(text: str) -> str:
    return re.sub(r"\s+", "", text).strip("：:").lower()


def is_toc_heading(line: str) -> bool:
    match = re.match(r"^#{1,6}\s+(.+?)\s*$", line)
    return bool(match and normalized(match.group(1)) in {"目录", "contents", "tableofcontents"})


def preprocess_markdown(input_path: Path, output_path: Path) -> bool:
    """Make # the document title and map Markdown ##/###/#### to Word H1/H2/H3."""
    lines = input_path.read_text(encoding="utf-8").splitlines()
    title = None
    subtitle = None
    body: list[str] = []
    seen_body = False
    has_toc = False

    for line in lines:
        top_title = heading_text(line, 1)
        if top_title and not seen_body and title is None:
            title = top_title
            continue
        if top_title and not seen_body and subtitle is None:
            subtitle = top_title
            continue
        if is_toc_heading(line):
            has_toc = True
            continue

        if line.strip():
            seen_body = True

        heading = re.match(r"^(#{2,6})(\s+.+)$", line)
        if heading:
            hashes, rest = heading.groups()
            body.append(hashes[1:] + rest)
        else:
            body.append(line)

    header = ["---"]
    if title:
        header.append(f'title: "{escape_yaml(title)}"')
    if subtitle:
        header.append(f'subtitle: "{escape_yaml(subtitle)}"')
    header.extend(["---", ""])
    output_path.write_text("\n".join(header + body) + "\n", encoding="utf-8")
    return has_toc


def set_run_font(run, size_pt=12, cn_font=FONT_CN, en_font=FONT_EN, bold=False):
    run.font.name = en_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = BLACK


def char_width(chars: float, size_pt: float = 12):
    return Pt(chars * size_pt)


def set_para(paragraph, align=None, before=0, after=0, line=None, first=None, hanging=None):
    fmt = paragraph.paragraph_format
    if align is not None:
        paragraph.alignment = align
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line is None:
        fmt.line_spacing = 1
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        fmt.line_spacing = Pt(line)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.first_line_indent = first
    if hanging is not None:
        fmt.left_indent = hanging
        fmt.first_line_indent = -hanging


def apply_font(paragraph, size=12, cn_font=FONT_CN, en_font=FONT_EN, bold=False):
    for run in paragraph.runs:
        set_run_font(run, size, cn_font, en_font, bold)


def configure_style(doc: Document, name: str, *, size=12, cn_font=FONT_CN, en_font=FONT_EN,
                    bold=False, before=0, after=0, line=22, align=None):
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = en_font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = BLACK
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line is None:
        fmt.line_spacing = 1
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        fmt.line_spacing = Pt(line)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if align is not None:
        fmt.alignment = align


def create_reference_docx(path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    configure_style(doc, "Normal", size=12, cn_font=FONT_CN, en_font=FONT_EN, line=22)
    configure_style(doc, "Body Text", size=12, cn_font=FONT_CN, en_font=FONT_EN, line=22)
    configure_style(doc, "Title", size=22, cn_font=FONT_CN, en_font=FONT_EN, bold=True,
                    line=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    configure_style(doc, "Subtitle", size=18, cn_font=FONT_CN, en_font=FONT_EN, bold=True,
                    line=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    configure_style(doc, "Heading 1", size=16, cn_font=FONT_CN_HEADING, en_font=FONT_EN,
                    bold=True, before=12, after=12, line=None, align=WD_ALIGN_PARAGRAPH.CENTER)
    configure_style(doc, "Heading 2", size=12, cn_font=FONT_CN_HEADING, en_font=FONT_EN,
                    bold=True, before=6, after=6, line=None, align=WD_ALIGN_PARAGRAPH.LEFT)
    configure_style(doc, "Heading 3", size=12, cn_font=FONT_CN, en_font=FONT_EN,
                    bold=True, before=6, after=6, line=None, align=WD_ALIGN_PARAGRAPH.LEFT)
    configure_style(doc, "Caption", size=10.5, cn_font=FONT_CN, en_font=FONT_EN,
                    line=None, before=6, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    configure_style(doc, "TOC Heading", size=18, cn_font=FONT_CN, en_font=FONT_EN,
                    bold=True, line=None, before=6, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    for name in ("TOC 1", "TOC 2", "TOC 3"):
        configure_style(doc, name, size=12, cn_font=FONT_CN, en_font=FONT_EN,
                        line=22, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_paragraph("Reference document for Pandoc styles.")
    doc.save(path)


def paragraph_kind(paragraph, mode: str) -> tuple[str, str]:
    text = paragraph.text.strip()
    norm = normalized(text)
    style = paragraph.style.name if paragraph.style else ""

    if style == "Title":
        return "title_cn", mode
    if style == "Subtitle":
        return "title_en", mode
    if norm in {"摘要", "中文摘要"}:
        return "abstract_title_cn", "abstract_cn"
    if norm in {"abstract", "英文摘要"}:
        return "abstract_title_en", "abstract_en"
    if norm in {"目录", "contents", "tableofcontents"} or style == "TOC Heading":
        return "toc_title", "body"
    if norm in {"参考文献", "references"}:
        return "heading1", "refs"
    if re.match(r"^(关键词|关键字)[:：]", text):
        return "keywords_cn", mode
    if re.match(r"^(keywords|key words)\s*[:：]", text, re.I):
        return "keywords_en", mode
    if re.match(r"^(图|figure|表|table)\s*[\d一二三四五六七八九十]+", text, re.I):
        return "caption", mode
    if style.startswith("TOC "):
        return "toc_entry", mode
    if style == "Heading 1":
        return "heading1", "body"
    if style == "Heading 2":
        return "heading2", "body"
    if style == "Heading 3":
        return "heading3", "body"
    if mode == "abstract_cn":
        return "abstract_body_cn", mode
    if mode == "abstract_en":
        return "abstract_body_en", mode
    if mode == "refs":
        return "reference", mode
    return "body", mode


def insert_page_breaks_between_heading1_sections(doc: Document):
    heading_indices = [
        idx for idx, paragraph in enumerate(doc.paragraphs)
        if paragraph.style and paragraph.style.name == "Heading 1"
    ]
    for idx in reversed(heading_indices[1:]):
        previous = doc.paragraphs[idx - 1]
        previous.add_run().add_break(WD_BREAK.PAGE)


def clear_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_table_edge(table, edge: str, value="single", size="8", color="000000"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    element = borders.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        borders.append(element)
    element.set(qn("w:val"), value)
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), color)


def set_cell_border(cell, edge: str, value="single", size="8", color="000000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    element = borders.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        borders.append(element)
    element.set(qn("w:val"), value)
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), color)


def set_cell_margins(cell, margin_twips="36"):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge in ("top", "left", "bottom", "right"):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), margin_twips)
        element.set(qn("w:type"), "dxa")


def apply_three_line_table(table):
    clear_table_borders(table)
    set_table_edge(table, "top", size="12")
    set_table_edge(table, "bottom", size="12")
    if table.rows:
        for cell in table.rows[0].cells:
            set_cell_border(cell, "bottom", size="8")
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, "36")
            for paragraph in cell.paragraphs:
                set_para(paragraph, before=0, after=0, line=None)
                apply_font(paragraph, 10.5, FONT_CN, FONT_EN, False)


def apply_chinese_article_format(docx_path: Path):
    doc = Document(docx_path)
    mode = "body"
    insert_page_breaks_between_heading1_sections(doc)

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        kind, mode = paragraph_kind(paragraph, mode)
        if kind == "title_cn":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, line=None)
            apply_font(paragraph, 22, FONT_CN, FONT_EN, True)
        elif kind == "title_en":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, line=None)
            apply_font(paragraph, 18, FONT_CN, FONT_EN, True)
        elif kind == "abstract_title_cn":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=18, line=None)
            apply_font(paragraph, 18, FONT_CN, FONT_EN, True)
        elif kind == "abstract_title_en":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=18, line=None)
            apply_font(paragraph, 18, FONT_CN, FONT_EN, True)
        elif kind == "keywords_cn":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=0, line=22)
            apply_font(paragraph, 12, FONT_CN, FONT_EN, False)
            if paragraph.runs:
                set_run_font(paragraph.runs[0], 12, FONT_CN, FONT_EN, True)
        elif kind == "keywords_en":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=0, line=22)
            apply_font(paragraph, 12, FONT_CN, FONT_EN, False)
            if paragraph.runs:
                set_run_font(paragraph.runs[0], 12, FONT_CN, FONT_EN, True)
        elif kind == "toc_title":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=18, line=None)
            apply_font(paragraph, 18, FONT_CN, FONT_EN, True)
        elif kind == "toc_entry":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=22)
            apply_font(paragraph, 12, FONT_CN, FONT_EN, False)
        elif kind == "heading1":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=12, line=None)
            apply_font(paragraph, 16, FONT_CN_HEADING, FONT_EN, True)
        elif kind == "heading2":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=6, line=None)
            apply_font(paragraph, 12, FONT_CN_HEADING, FONT_EN, True)
        elif kind == "heading3":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=6, line=None)
            apply_font(paragraph, 12, FONT_CN, FONT_EN, True)
        elif kind == "caption":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6, line=None)
            apply_font(paragraph, 10.5, FONT_CN, FONT_EN, False)
        elif kind == "reference":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=20,
                     hanging=char_width(2.5, 10.5))
            apply_font(paragraph, 10.5, FONT_CN, FONT_EN, False)
        elif kind == "abstract_body_en":
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=22,
                     first=char_width(1, 12))
            apply_font(paragraph, 12, FONT_CN, FONT_EN, False)
        else:
            set_para(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=22,
                     first=char_width(2, 12))
            apply_font(paragraph, 12, FONT_CN, FONT_EN, False)

    for table in doc.tables:
        apply_three_line_table(table)

    doc.save(docx_path)


def convert(markdown_path: Path, output_path: Path, *, install_if_missing: bool = True):
    pandoc = find_pandoc()
    if not pandoc:
        if not install_if_missing:
            raise RuntimeError("Pandoc is missing. Re-run without --no-install-pandoc to install it.")
        pandoc = install_pandoc()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        reference_doc = tmp_dir / "reference.docx"
        preprocessed_md = tmp_dir / "input.md"
        create_reference_docx(reference_doc)
        has_toc = preprocess_markdown(markdown_path, preprocessed_md)
        cmd = [
            pandoc,
            str(preprocessed_md),
            "--from",
            "markdown+pipe_tables+implicit_figures",
            "--to",
            "docx",
            "--standalone",
            "--reference-doc",
            str(reference_doc),
            "-o",
            str(output_path),
        ]
        if has_toc:
            cmd.extend(["--toc", "--toc-depth=3"])
        run(cmd)
    apply_chinese_article_format(output_path)


def main():
    parser = argparse.ArgumentParser(description="Convert Markdown to standardized DOCX with Pandoc.")
    parser.add_argument("input", type=Path, help="Input UTF-8 Markdown file.")
    parser.add_argument("output", type=Path, help="Output DOCX file.")
    parser.add_argument("--no-install-pandoc", action="store_true",
                        help="Fail instead of installing Pandoc when it is missing.")
    args = parser.parse_args()
    convert(args.input, args.output, install_if_missing=not args.no_install_pandoc)


if __name__ == "__main__":
    main()
